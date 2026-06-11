from __future__ import annotations

import json
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


DOCS_DIR = Path(__file__).resolve().parent
VISUALS_DIR = DOCS_DIR / "visuals"
GENERATED_DIR = VISUALS_DIR / "generated"
OUT = DOCS_DIR / "VRBS_LocateAnything_Result_Archive_with_visuals.docx"


def load_results() -> dict:
    return json.loads((VISUALS_DIR / "text" / "results.json").read_text(encoding="utf-8"))


def load_category_rows() -> list[dict]:
    rows: list[dict] = []
    path = VISUALS_DIR / "text" / "category_predictions.jsonl"
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            rows.append(json.loads(line))
    rows.sort(key=lambda row: row["query_idx"])
    return rows


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = load_font(32, bold=True)
FONT_SUBTITLE = load_font(24, bold=True)
FONT_BODY = load_font(21)
FONT_SMALL = load_font(18)
FONT_TINY = load_font(15)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    if not text:
        return 0, 0
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = str(text).replace("\n", " ").split()
    if not words:
        return [""]
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if text_size(draw, candidate, font)[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    max_width: int,
    line_gap: int = 7,
    max_lines: int | None = None,
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, font, max_width)
    if max_lines is not None and len(lines) > max_lines:
        lines = lines[:max_lines]
        if lines:
            lines[-1] = lines[-1].rstrip(".") + "..."
    line_height = text_size(draw, "Ag", font)[1] + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def fit_image(path: Path, width: int, height: int, fill: str = "#FFFFFF") -> Image.Image:
    canvas = Image.new("RGB", (width, height), fill)
    if not path.exists():
        draw = ImageDraw.Draw(canvas)
        draw.rectangle((0, 0, width - 1, height - 1), outline="#DADCE0", width=2)
        draw_wrapped(draw, (18, 18), f"Missing image: {path.name}", FONT_SMALL, "#9AA0A6", width - 36)
        return canvas
    with Image.open(path) as raw:
        img = ImageOps.exif_transpose(raw).convert("RGB")
    img.thumbnail((width, height), Image.Resampling.LANCZOS)
    x = (width - img.width) // 2
    y = (height - img.height) // 2
    canvas.paste(img, (x, y))
    return canvas


def frame_image(img: Image.Image, title: str, caption: str, width: int, height: int) -> Image.Image:
    tile = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(tile)
    draw.rounded_rectangle((0, 0, width - 1, height - 1), radius=12, outline="#DADCE0", width=2)
    draw.text((18, 14), title, font=FONT_SUBTITLE, fill="#202124")
    image_top = 56
    image_h = height - 116
    inner = fit_image(Path("__already_loaded__"), width - 36, image_h)
    inner.paste(img.resize(inner.size) if False else img, (0, 0))
    tile.paste(img, (18 + (width - 36 - img.width) // 2, image_top + (image_h - img.height) // 2))
    draw_wrapped(draw, (18, height - 54), caption, FONT_TINY, "#5F6368", width - 36, max_lines=2)
    return tile


def make_input_tile(result: dict, tile_w: int = 760, tile_h: int = 360) -> Image.Image:
    q = result["query_idx"]
    tile = Image.new("RGB", (tile_w, tile_h), "#FFFFFF")
    draw = ImageDraw.Draw(tile)
    draw.rounded_rectangle((0, 0, tile_w - 1, tile_h - 1), radius=12, outline="#DADCE0", width=2)
    draw.text((18, 14), f"q{q}: input pair", font=FONT_SUBTITLE, fill="#202124")
    draw_wrapped(draw, (18, 48), result["question"], FONT_SMALL, "#3C4043", tile_w - 36, max_lines=2)

    human_path = VISUALS_DIR / "inputs" / f"q{q}" / f"ask_human_{q}.png"
    scene_path = VISUALS_DIR / "inputs" / f"q{q}" / f"ask_bg_{q}.jpg"
    human = fit_image(human_path, 210, 210, "#F8F9FA")
    scene = fit_image(scene_path, 480, 210, "#F8F9FA")
    tile.paste(human, (24, 116))
    tile.paste(scene, (254, 116))
    draw.text((24, 328), "Image A: user photo", font=FONT_TINY, fill="#5F6368")
    draw.text((254, 328), "Image B: test scene", font=FONT_TINY, fill="#5F6368")
    return tile


def make_image_tile(path: Path, title: str, caption: str, tile_w: int = 760, tile_h: int = 500) -> Image.Image:
    tile = Image.new("RGB", (tile_w, tile_h), "#FFFFFF")
    draw = ImageDraw.Draw(tile)
    draw.rounded_rectangle((0, 0, tile_w - 1, tile_h - 1), radius=12, outline="#DADCE0", width=2)
    draw.text((18, 14), title, font=FONT_SUBTITLE, fill="#202124")
    img = fit_image(path, tile_w - 36, tile_h - 120, "#F8F9FA")
    tile.paste(img, (18, 58))
    draw_wrapped(draw, (18, tile_h - 52), caption, FONT_TINY, "#5F6368", tile_w - 36, max_lines=2)
    return tile


def make_grid(name: str, tiles: list[Image.Image], cols: int = 2, gap: int = 24) -> Path:
    if not tiles:
        raise ValueError("no tiles supplied")
    rows = (len(tiles) + cols - 1) // cols
    tile_w, tile_h = tiles[0].size
    canvas = Image.new(
        "RGB",
        (cols * tile_w + (cols + 1) * gap, rows * tile_h + (rows + 1) * gap),
        "#FFFFFF",
    )
    for idx, tile in enumerate(tiles):
        row, col = divmod(idx, cols)
        x = gap + col * (tile_w + gap)
        y = gap + row * (tile_h + gap)
        canvas.paste(tile, (x, y))
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    out = GENERATED_DIR / name
    canvas.save(out, quality=92)
    return out


def make_category_table(rows: list[dict]) -> Path:
    width = 1600
    row_h = 72
    height = 130 + row_h * len(rows)
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((32, 24), "Step 2 result: Qwen3.5 category / grounding phrase extraction", font=FONT_TITLE, fill="#202124")
    headers = ["Query", "User command", "Category", "LocateAnything prompt"]
    x_positions = [32, 160, 720, 1010]
    col_widths = [110, 520, 250, 520]
    y = 86
    draw.rectangle((24, y, width - 24, y + 42), outline="#DADCE0", fill="#F8F9FA")
    for x, header in zip(x_positions, headers):
        draw.text((x, y + 10), header, font=FONT_SMALL, fill="#202124")
    y += 42
    for idx, row in enumerate(rows):
        fill = "#FFFFFF" if idx % 2 == 0 else "#FAFAFA"
        draw.rectangle((24, y, width - 24, y + row_h), outline="#E0E0E0", fill=fill)
        values = [
            f"q{row['query_idx']}",
            row["question"],
            row["category"],
            row["prompt"],
        ]
        for x, col_w, value in zip(x_positions, col_widths, values):
            draw_wrapped(draw, (x, y + 12), value, FONT_SMALL, "#3C4043", col_w, max_lines=2)
        y += row_h
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    out = GENERATED_DIR / "step2_category_extraction.jpg"
    img.save(out, quality=92)
    return out


def make_video_memory_card() -> Path:
    text_path = VISUALS_DIR / "text" / "phase1_video_context.txt"
    memory = text_path.read_text(encoding="utf-8").strip()
    words = memory.split()
    width, height = 1600, 760
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((32, 28), "Step 4 result: first-round video understanding memory", font=FONT_TITLE, fill="#202124")
    draw.text((32, 78), f"{len(memory)} characters | {len(words)} words | reused as isolated base context for q1-q10", font=FONT_SMALL, fill="#5F6368")
    draw.rounded_rectangle((32, 128, width - 32, height - 32), radius=16, outline="#DADCE0", width=2, fill="#F8F9FA")
    excerpt = " ".join(memory.split())
    y = draw_wrapped(
        draw,
        (60, 158),
        excerpt,
        FONT_BODY,
        "#202124",
        width - 120,
        line_gap=10,
        max_lines=17,
    )
    if y < height - 76:
        draw.text((60, y + 8), "[excerpt truncated in figure; full text is embedded in docs/visuals/text/phase1_video_context.txt]", font=FONT_TINY, fill="#5F6368")
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    out = GENERATED_DIR / "step4_video_memory_card.jpg"
    img.save(out, quality=92)
    return out


def make_per_query_panel(result: dict) -> Path:
    q = result["query_idx"]
    width, height = 1700, 600
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = (
        f"q{q}: {result['question']} | prompt: {result['category_prompt']} | "
        f"tag: {result['selection']['selected_tag']} | IoU: {result['max_iou']:.4f} | "
        f"{'HIT' if result['hit'] else 'MISS'}"
    )
    draw.text((32, 24), title, font=FONT_SUBTITLE, fill="#202124")
    human = fit_image(VISUALS_DIR / "inputs" / f"q{q}" / f"ask_human_{q}.png", 260, 390, "#F8F9FA")
    tagged = fit_image(VISUALS_DIR / "output_images" / f"q{q}_tagged.jpg", 640, 390, "#F8F9FA")
    eval_name = Path(result["eval_image"]).name
    final = fit_image(VISUALS_DIR / "output_images" / eval_name, 640, 390, "#F8F9FA")
    blocks = [
        (32, 94, 260, human, "Image A: user photo"),
        (330, 94, 640, tagged, "Tagged Image B: LocateAnything proposals"),
        (1030, 94, 640, final, "Final eval: selected bbox vs GT"),
    ]
    for x, y, w, panel, caption in blocks:
        draw.rounded_rectangle((x - 8, y - 8, x + w + 8, y + 438), radius=12, outline="#DADCE0", width=2)
        img.paste(panel, (x, y))
        draw_wrapped(draw, (x, y + 404), caption, FONT_SMALL, "#5F6368", w, max_lines=1)
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    out = GENERATED_DIR / f"query_q{q:02d}_panel.jpg"
    img.save(out, quality=92)
    return out


def generate_figures(results: dict) -> dict[str, list[Path] | Path]:
    rows = results["results"]
    categories = load_category_rows()

    input_paths = [
        make_grid(
            "step1_inputs_q1_q5.jpg",
            [make_input_tile(row) for row in rows[:5]],
        ),
        make_grid(
            "step1_inputs_q6_q10.jpg",
            [make_input_tile(row) for row in rows[5:]],
        ),
    ]
    category_path = make_category_table(categories)
    video_memory_path = make_video_memory_card()

    tagged_paths: list[Path] = []
    eval_paths: list[Path] = []
    for start, end in [(0, 5), (5, 10)]:
        tagged_tiles = []
        eval_tiles = []
        for row in rows[start:end]:
            q = row["query_idx"]
            tagged_tiles.append(
                make_image_tile(
                    VISUALS_DIR / "output_images" / f"q{q}_tagged.jpg",
                    f"q{q}: proposal tags",
                    f"{len(row['proposals'])} LocateAnything proposal(s); VLM selected tag {row['selection']['selected_tag']}",
                )
            )
            eval_tiles.append(
                make_image_tile(
                    VISUALS_DIR / "output_images" / Path(row["eval_image"]).name,
                    f"q{q}: final bbox evaluation",
                    f"IoU {row['max_iou']:.4f}; {'HIT' if row['hit'] else 'MISS'} at threshold {results['iou_threshold']}",
                )
            )
        tagged_paths.append(make_grid(f"step3_tagged_q{start + 1}_q{end}.jpg", tagged_tiles))
        eval_paths.append(make_grid(f"step6_eval_q{start + 1}_q{end}.jpg", eval_tiles))

    panels = [make_per_query_panel(row) for row in rows]
    return {
        "input_grids": input_paths,
        "category": category_path,
        "video_memory": video_memory_path,
        "tagged_grids": tagged_paths,
        "eval_grids": eval_paths,
        "query_panels": panels,
    }


def set_cell_border(cell, color: str = "DADCE0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_run(run, size=11, bold=False, color="000000") -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=20 if level == 1 else 18, after=6)
    run = paragraph.add_run(text)
    style_run(run, size=20 if level == 1 else 16, bold=False, color="000000")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph)
    run = paragraph.add_run(text)
    style_run(run)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    style_run(run)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=2, after=10)
    run = paragraph.add_run(text)
    style_run(run, size=9, color="555555")


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=2)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        para = hdr[idx].paragraphs[0]
        set_paragraph_spacing(para, after=0)
        run = para.add_run(text)
        style_run(run, bold=True)
        set_cell_border(hdr[idx])
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            para = cells[idx].paragraphs[0]
            set_paragraph_spacing(para, after=0)
            run = para.add_run(text)
            style_run(run, size=10)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build() -> None:
    results = load_results()
    figures = generate_figures(results)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("VRBS LocateAnything Pipeline Result Archive With Visuals")
    style_run(title_run, size=26, bold=False)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12)
    subtitle_run = subtitle.add_run(
        "Project: /mnt/si00068187c7/default/myc/projects/vrbs_locateanything | "
        "Final run: dengnan_locateanything_all_full_qwen35_prompt0606_contextcopy"
    )
    style_run(subtitle_run, size=10, color="555555")

    add_heading(doc, "Summary")
    add_body(
        doc,
        "This image-rich archive documents the final VRBS LocateAnything pipeline. It includes "
        "step-level visual outputs, per-query input and prediction panels, and the final evaluation "
        "metrics from the 10-question full test run.",
    )
    add_bullet(doc, f"Final precision at IoU {results['iou_threshold']}: {results['precision'] * 100:.2f}% ({results['hits']} correct out of {results['total']} queries).")
    add_bullet(doc, "The pipeline uses Qwen/Qwen3.5-0.8B for isolated category extraction, LocateAnything for 2D proposals, and the distilled Qwen3.6 VLM for tag selection.")
    add_bullet(doc, "Each query copies only the first-round video memory as base context; previous second-round question/answer state is not reused.")

    add_heading(doc, "Step Visuals")
    add_body(doc, "Step 1 shows the user-person image and the scene image supplied to each query.")
    for idx, path in enumerate(figures["input_grids"], start=1):
        add_picture(doc, path, f"Figure 1.{idx}. Step 1 inputs: user photo and target scene pairs.")

    add_body(doc, "Step 2 shows the small language model output used as the LocateAnything grounding phrase.")
    add_picture(doc, figures["category"], "Figure 2. Step 2 category extraction and grounding phrase results.")

    add_body(doc, "Step 3 shows LocateAnything 2D proposals rendered as numbered tags on the test scene image.")
    for idx, path in enumerate(figures["tagged_grids"], start=1):
        add_picture(doc, path, f"Figure 3.{idx}. Step 3 tagged proposal images.")

    add_body(doc, "Step 4 shows the first-round video understanding memory. The full memory is kept in the local text artifact and copied as isolated base context for every query.")
    add_picture(doc, figures["video_memory"], "Figure 4. Step 4 video memory summary and excerpt.")

    add_body(doc, "Step 5 and Step 6 are shown together below: the VLM-selected tag is converted back to a 2D bbox and evaluated against the ground-truth box.")
    for idx, path in enumerate(figures["eval_grids"], start=1):
        add_picture(doc, path, f"Figure 5.{idx}. Step 5-6 selected bbox evaluation images.")

    add_heading(doc, "Per-Query Results")
    result_rows = []
    for row in results["results"]:
        result_rows.append(
            [
                f"q{row['query_idx']}",
                row["question"],
                row["category_prompt"],
                str(len(row["proposals"])),
                str(row["selection"]["selected_tag"]),
                f"{row['max_iou']:.4f}",
                "HIT" if row["hit"] else "MISS",
            ]
        )
    add_table(
        doc,
        ["Query", "User command", "Grounding phrase", "Props", "Tag", "IoU", "Result"],
        result_rows,
        [0.45, 2.05, 1.35, 0.45, 0.45, 0.55, 0.70],
    )

    add_heading(doc, "Per-Query Visual Appendix")
    add_body(
        doc,
        "Each panel shows the query's Image A user photo, the tagged Image B proposal image, and the final evaluation visualization."
    )
    for path in figures["query_panels"]:
        q_name = path.stem.replace("query_", "").replace("_panel", "")
        add_picture(doc, path, f"Figure A.{q_name}. Per-query input, proposal, and final bbox result.", width=6.5)

    add_heading(doc, "Artifacts")
    artifact_rows = [
        ["Project code", "/mnt/si00068187c7/default/myc/projects/vrbs_locateanything"],
        ["Final output directory", "/mnt/si00068187c7/default/myc/data/VRBS/output/dengnan_locateanything_all_full_qwen35_prompt0606_contextcopy"],
        ["Downloaded local visual assets", "docs/visuals"],
        ["Generated composite figures", "docs/visuals/generated"],
        ["Main entrypoint", "run_pipeline.py"],
        ["Prompt definitions", "src/prompts.py"],
        ["Result summary", "docs/visuals/text/results.json"],
    ]
    add_table(doc, ["Artifact", "Path / file"], artifact_rows, [1.8, 4.7])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
