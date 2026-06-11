from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "VRBS_LocateAnything_Result_Archive.docx"


QUESTIONS = {
    1: "Where are my white headphones?",
    2: "Help me get my cup.",
    3: "Help me find my laptop.",
    4: "Where did I put my headphones?",
    5: "Help me find where my phone is.",
    6: "Where is my game controller?",
    7: "Help me find where my headphones are placed.",
    8: "Where did my game controller go?",
    9: "Help me find my gaming laptop.",
    10: "Where can I find my white powercube?",
}

RESULTS = [
    (1, "white headphones", 1, 1, 0.9606, "HIT"),
    (2, "cup", 1, 1, 0.9627, "HIT"),
    (3, "laptop", 1, 1, 0.9464, "HIT"),
    (4, "headphones", 3, 2, 0.9585, "HIT"),
    (5, "phone", 1, 1, 0.7925, "HIT"),
    (6, "game controller", 2, 1, 0.9528, "HIT"),
    (7, "headphones", 3, 3, 0.9587, "HIT"),
    (8, "game controller", 2, 2, 0.6007, "HIT"),
    (9, "gaming laptop", 1, 1, 0.9811, "HIT"),
    (10, "white power cube", 1, 1, 0.9200, "HIT"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DADCE0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_run(run, size=11, bold=False, color="000000") -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=20 if level == 1 else 18, after=6)
    run = paragraph.add_run(text)
    style_run(run, size=20 if level == 1 else 16, bold=False, color="000000")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph)
    run = paragraph.add_run(text)
    style_run(run)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    style_run(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        para = hdr[idx].paragraphs[0]
        set_paragraph_spacing(para, after=0)
        run = para.add_run(text)
        style_run(run, bold=True)
        set_cell_border(hdr[idx])
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            para = cells[idx].paragraphs[0]
            set_paragraph_spacing(para, after=0)
            run = para.add_run(text)
            style_run(run, size=10)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("VRBS LocateAnything Pipeline Result Archive")
    style_run(title_run, size=26, bold=False)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12)
    subtitle_run = subtitle.add_run(
        "Project: /mnt/si00068187c7/default/myc/projects/vrbs_locateanything | "
        "Final run: dengnan_locateanything_all_full_qwen35_prompt0606_contextcopy"
    )
    style_run(subtitle_run, size=10, color="555555")

    add_heading(doc, "Summary")
    add_body(
        doc,
        "The final evaluated pipeline uses Qwen/Qwen3.5-0.8B to extract a grounding phrase, "
        "LocateAnything to generate 2D proposals, and the distilled Qwen3.6 VLM to select the "
        "correct proposal tag using a two-round prompt structure adapted from the 0606 VRBS script.",
    )
    add_bullet(doc, "Final precision at IoU 0.5: 100.00% (10 correct out of 10 queries).")
    add_bullet(doc, "No selected tag in the final run used the forced fallback path; all tags were parsed from VLM outputs.")
    add_bullet(doc, "The first-phase video memory was reused for every query, but each second-phase query was isolated from previous query/answer context.")

    add_heading(doc, "Pipeline Steps")
    step_rows = [
        [
            "1. Data inputs",
            "Video: data/VRBS/input/dengnan/video/video.mp4; per-query ask_*.txt, ask_bg_*.jpg, ask_human_*.png; GT: gts/items_annotations.json.",
            "10 query records discovered and paired with GT boxes.",
        ],
        [
            "2. Category extraction",
            "User command text for each query; model: /mnt/si00068187c7/default/myc/models/Qwen3.5-0.8B.",
            "Short LocateAnything grounding phrases, with explicit attribute filtering to prevent hallucinated colors or brands.",
        ],
        [
            "3. 2D proposal generation",
            "Scene image ask_bg_*.jpg plus Qwen3.5 grounding phrase; model: nvidia/LocateAnything-3B.",
            "1-3 candidate boxes per query, rendered as numbered tags on the scene image.",
        ],
        [
            "4. Video memory",
            "Original video plus phase1 prompt adapted from the VRBS 0606 script; VLM: Qwen3.6-35B-A3B.",
            "A 6157-character, 1054-word video memory describing people, interactions, item attributes, and final locations.",
        ],
        [
            "5. Tag selection",
            "Image A, tagged Image B, user command, candidate tags, and first-phase video memory.",
            "A selected_tag JSON value, using the 0606-style chain: match person, parse target item, retrieve interacted item, choose tag.",
        ],
        [
            "6. Evaluation",
            "Selected proposal bbox and GT bbox; IoU threshold = 0.5.",
            "10/10 hits, precision = 100.00%.",
        ],
    ]
    add_table(doc, ["Step", "Input", "Output"], step_rows, [1.35, 3.2, 1.95])

    add_heading(doc, "Per-Query Results")
    result_rows = []
    for q, prompt, proposal_count, selected_tag, iou, status in RESULTS:
        result_rows.append(
            [
                f"q{q}",
                QUESTIONS[q],
                prompt,
                str(proposal_count),
                str(selected_tag),
                f"{iou:.4f}",
                status,
            ]
        )
    add_table(
        doc,
        ["Query", "User command", "Grounding phrase", "Props", "Tag", "IoU", "Result"],
        result_rows,
        [0.45, 2.05, 1.35, 0.45, 0.45, 0.55, 0.70],
    )

    add_heading(doc, "Prompt And Context Handling")
    add_body(
        doc,
        "The second-round prompt now follows the original 0606 localization prompt structure, but asks the VLM to select a numbered tag instead of regressing absolute coordinates. The retained reasoning sequence is: match the person in Image A by facial features, parse the user command, identify the matching interacted item from the video memory, and select only among LocateAnything-generated candidate tags.",
    )
    add_body(
        doc,
        "For context isolation, the implementation stores the phase1 chain as system + video user + assistant video memory. For every q1-q10 request it copies only that base history and appends the current query. The current query and answer are never written back, so later questions cannot inherit prior question context.",
    )

    add_heading(doc, "Final Artifacts")
    artifact_rows = [
        ["Project code", "/mnt/si00068187c7/default/myc/projects/vrbs_locateanything"],
        ["Final output directory", "/mnt/si00068187c7/default/myc/data/VRBS/output/dengnan_locateanything_all_full_qwen35_prompt0606_contextcopy"],
        ["Main entrypoint", "run_pipeline.py"],
        ["Prompt definitions", "src/prompts.py"],
        ["VLM tag selector", "src/vlm_selector.py"],
        ["Category extractor", "src/category_extractor.py and src/category_extractor_worker.py"],
        ["Result summary", "text/results.json"],
        ["Tagged/evaluation images", "images/q*_tagged.jpg and images/vis_q*.jpg"],
    ]
    add_table(doc, ["Artifact", "Path / file"], artifact_rows, [1.8, 4.7])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
